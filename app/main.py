import base64
import io
import logging
from typing import Optional

from fastapi import FastAPI
from fastapi.responses import JSONResponse
from pydantic import ValidationError

from .models import RenderDocxRequest, RenderXlsxRequest, RenderResponse

from docx import Document
from html2docx import html2docx
from openpyxl import Workbook

logger = logging.getLogger("document-renderer")
logging.basicConfig(level=logging.INFO)

app = FastAPI(title="Document Render Service", version="1.0.0")


@app.get("/health")
def health() -> dict:
    return {"status": "ok"}


def _build_error_response(message: str) -> RenderResponse:
    logger.error(f"Render error: {message}")
    return RenderResponse(data=None, error=True, message=message)


@app.post("/render-docx", response_model=RenderResponse)
def render_docx(req: RenderDocxRequest):
    """
    Render HTML to DOCX and return as base64.
    """
    try:
        if not req.html or not req.title:
            return _build_error_response("Missing html or title")

        # --- Hauptpfad: html2docx korrekt verwenden ---
        try:
            # html2docx gibt ein io.BytesIO-Objekt zurück
            buf = html2docx(req.html, title=req.title)
            # Sicherheitshalber auf Anfang setzen
            buf.seek(0)
            raw_bytes = buf.read()

        except Exception as e:
            logger.warning(
                f"html2docx failed, falling back to simple document. Error: {e}"
            )
            # Fallback: ganz simples Dokument mit Hinweis
            doc = Document()
            doc.add_heading(req.title, level=1)
            doc.add_paragraph(
                "HTML could not be fully parsed. Fallback content only."
            )

            tmp = io.BytesIO()
            doc.save(tmp)
            tmp.seek(0)
            raw_bytes = tmp.read()

        # Gemeinsamer Pfad: Bytes → Base64
        b64_data = base64.b64encode(raw_bytes).decode("utf-8")
        return RenderResponse(data=b64_data, error=False, message=None)

    except Exception as e:
        logger.exception("Unexpected error in /render-docx")
        return _build_error_response(f"Unexpected error while rendering DOCX: {e}")


@app.post("/render-xlsx", response_model=RenderResponse)
def render_xlsx(req: RenderXlsxRequest):
    """
    Render sheets JSON to XLSX and return as base64.
    """
    try:
        if not req.sheets or len(req.sheets) == 0:
            return _build_error_response("At least one sheet is required")

        wb = Workbook()
        # remove default sheet
        default_sheet = wb.active
        wb.remove(default_sheet)

        for sheet_def in req.sheets:
            # Excel sheet name max length 31
            ws = wb.create_sheet(title=sheet_def.name[:31])

            # Write headers
            if sheet_def.headers:
                ws.append(sheet_def.headers)

            # Write rows
            for row in sheet_def.rows:
                ws.append(row)

        buffer = io.BytesIO()
        wb.save(buffer)
        buffer.seek(0)

        b64_data = base64.b64encode(buffer.read()).decode("utf-8")
        return RenderResponse(data=b64_data, error=False, message=None)

    except Exception as e:
        logger.exception("Unexpected error in /render-xlsx")
        return _build_error_response(f"Unexpected error while rendering XLSX: {e}")


@app.exception_handler(ValidationError)
async def validation_exception_handler(request, exc: ValidationError):
    logger.warning(f"Validation error: {exc}")
    return JSONResponse(
        status_code=400,
        content=RenderResponse(
            data=None,
            error=True,
            message="Invalid request payload",
        ).dict(),
    )


@app.exception_handler(Exception)
async def general_exception_handler(request, exc: Exception):
    logger.exception("Unhandled error")
    return JSONResponse(
        status_code=500,
        content=RenderResponse(
            data=None,
            error=True,
            message="Internal server error",
        ).dict(),
    )
